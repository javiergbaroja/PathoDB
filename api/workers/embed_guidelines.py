"""
PathoDB — Guideline embedding ingestion (guideline RAG index builder, audit #5).

Embeds external reporting-guideline text (CAP protocols, ICCR datasets) into the
`guideline_chunks` pgvector table so the agent's `guideline_search` tool can
retrieve cited, version-stamped staging/grading/required-element standards.

Mirrors api/workers/embed_reports.py (buffer chunks across docs, embed in large
GPU batches, single multi-row INSERT per flush). Guideline-specific parts:
  - .docx text is extracted by walking the document body IN ORDER so TABLE
    content survives (CAP thresholds/checklists live in tables that
    doc.paragraphs drops); the nearest heading is tracked as the chunk `section`.
  - Metadata (organ, specimen, version, doc_slug) is parsed from the filename
    (api/agent/guideline_meta.py) plus inline title/date.
  - Idempotent + VERSION-AWARE: a doc whose content_hash is unchanged is skipped;
    a changed hash on the same doc_slug DELETEs the old rows and re-inserts, so a
    newer protocol version replaces the prior one ("latest only").

Run from the repo root (conda env `langchain`):
    python api/workers/embed_guidelines.py --source all
    python api/workers/embed_guidelines.py --source cap --limit 5

Requires: pgvector + schema applied (db/schema.sql: guideline_chunks),
python-docx, sentence-transformers, and settings.embedding_dim == model dim.
"""
import argparse
import hashlib
import logging
import sys
import time
from pathlib import Path

sys.path.insert(0, str(Path(__file__).resolve().parents[2]))

from sqlalchemy import text  # noqa: E402

from api.database import SessionLocal  # noqa: E402
from api.config import get_settings  # noqa: E402
from api.agent.embeddings import embed_texts  # noqa: E402
from api.agent.textutil import chunk_report  # noqa: E402
from api.agent.guideline_meta import (parse_guideline_filename, render_table_rows,  # noqa: E402
                                      organ_from_title, is_title_boilerplate,
                                      heading_level, element_rows_from_grid)

logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(message)s")
log = logging.getLogger("embed_guidelines")

import re  # noqa: E402

_POSTING_DATE_RE = re.compile(r"(?:Protocol\s+Posting\s+Date|Required\s+Use\s+Date)\s*:?\s*([^\n]{3,60})", re.I)


def _vector_literal(vec) -> str:
    return "[" + ",".join(f"{x:.6f}" for x in vec) + "]"


# ── docx extraction ──────────────────────────────────────────────────────────

def _all_bold(para) -> bool:
    runs = [r for r in para.runs if (r.text or "").strip()]
    return bool(runs) and all(bool(r.bold) for r in runs)


def _cell_text(cell) -> str:
    return " ".join(p.text for p in cell.paragraphs).strip()


# CAP protocols split into 3 zones by 13pt-bold headers. The reporting ELEMENTS
# live only in the template zone; the rest is boilerplate / explanatory prose.
_ZONE_KIND = {"frontmatter": "frontmatter", "template": "element", "notes": "note"}


def _iter_blocks(doc):
    """Yield (section, text, kind) blocks in document order, interleaving
    paragraphs and tables via the XML body so nothing is dropped.

    `kind` marks reporting ELEMENTS vs boilerplate/notes so enumeration lists only
    real elements. CAP: a zone tracker flips at 'Reporting Template' → 'Explanatory
    Notes', so template-zone content is 'element' and the lettered notes are
    'note'. ICCR: each reporting-table row is an 'element' (name as `section`,
    carrying core status/values); its definition/scope tables stay 'frontmatter'.
    Sections come from bold/ALL-CAPS headings (these docs use no Word heading
    styles): ALL-CAPS = category, bold = sub-element → 'CATEGORY — Element'."""
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    major, minor = "", ""     # ALL-CAPS section  →  bold sub-element
    section = ""
    zone = "frontmatter"
    for child in doc.element.body.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag == "p":
            para = Paragraph(child, doc)
            txt = (para.text or "").strip()
            if not txt:
                continue
            low = txt.lower()
            if low.startswith("reporting template"):
                zone = "template"
            elif low.startswith("explanatory note"):
                zone = "notes"
            lvl = heading_level(txt, _all_bold(para))
            if lvl == "major":
                major, minor = txt, ""             # new section resets the sub-element
            elif lvl == "minor":
                minor = txt
            if lvl:
                section = (f"{major} — {minor}" if major and minor else (major or minor))[:120]
            yield (section, txt, _ZONE_KIND[zone])
        elif tag == "tbl":
            table = Table(child, doc)
            grid = [[_cell_text(c) for c in row.cells] for row in table.rows]
            elements = element_rows_from_grid(grid)
            if elements:
                for name, text in elements:
                    yield (name[:120], text, "element")   # ICCR reporting element
            else:
                rendered = render_table_rows(grid)
                if rendered:
                    yield (section, rendered, _ZONE_KIND[zone])


def _extract(path: Path):
    """Return (blocks, full_text, title, doc_date) for one .docx, or None on error."""
    try:
        import docx
        doc = docx.Document(str(path))
    except Exception as e:
        log.warning("Skip %s: cannot open (%s)", path.name, e)
        return None
    blocks = list(_iter_blocks(doc))
    if not blocks:
        log.warning("Skip %s: no extractable text", path.name)
        return None
    full_text = "\n".join(t for _s, t, _k in blocks)
    # Title = first substantial NON-boilerplate paragraph (CAP 'Protocol for…',
    # ICCR '… Reporting Guide'), skipping CORE/NON-CORE legends + figure captions.
    title = next((t for _s, t, _k in blocks
                  if len(t) > 20 and not is_title_boilerplate(t)),
                 blocks[0][1])[:250]
    dm = _POSTING_DATE_RE.search(full_text)
    doc_date = dm.group(1).strip() if dm else None
    return blocks, full_text, title, doc_date


def _doc_chunks(blocks, max_chars, overlap):
    """Chunk the document per section (so each chunk keeps a clean `section` tag),
    preserving document order. Returns [(section, chunk_index, chunk_text, kind), …]
    with a document-global chunk_index (matches UNIQUE(doc_slug, chunk_index)).
    `kind` is taken from the section's first block (kind is constant within a
    section — a zone change always coincides with a heading, i.e. a new section)."""
    blocks = list(blocks)     # tolerate a generator
    out = []
    if not blocks:
        return out
    ci = 0
    cur_section, _, cur_kind = blocks[0]
    buf = []

    def emit(section, kind):
        nonlocal ci
        if not buf:
            return
        for ch in chunk_report("\n".join(buf), max_chars, overlap):
            out.append((section, ci, ch, kind))
            ci += 1

    for sec, txt, kind in blocks:
        if sec != cur_section:
            emit(cur_section, cur_kind)
            buf.clear()
            cur_section, cur_kind = sec, kind
        buf.append(txt)
    emit(cur_section, cur_kind)
    return out


# ── DB helpers ────────────────────────────────────────────────────────────────

def _bulk_insert(db, rows) -> None:
    """One multi-row INSERT. `rows`: (source_org, doc_slug, kind, title, organ,
    specimen_type, version, doc_date, section, chunk_index, chunk_text,
    vector_literal, source_path, content_hash)."""
    if not rows:
        return
    cols = ("source_org, doc_slug, kind, title, organ, specimen_type, version, "
            "doc_date, section, chunk_index, chunk_text, embedding, source_path, content_hash")
    try:
        from psycopg2.extras import execute_values
    except ImportError:
        execute_values = None
    if execute_values is not None:
        dbapi = db.connection().connection.dbapi_connection
        with dbapi.cursor() as cur:
            execute_values(
                cur,
                f"INSERT INTO guideline_chunks ({cols}) VALUES %s "
                "ON CONFLICT (doc_slug, chunk_index) DO NOTHING",
                rows,
                template="(%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s,%s::vector,%s,%s)",
                page_size=500,
            )
    else:
        db.execute(text(f"""
            INSERT INTO guideline_chunks ({cols})
            VALUES (:org,:slug,:kind,:title,:organ,:spec,:ver,:date,:sec,:ci,:ct,
                    CAST(:emb AS vector),:path,:hash)
            ON CONFLICT (doc_slug, chunk_index) DO NOTHING
        """), [{"org": r[0], "slug": r[1], "kind": r[2], "title": r[3], "organ": r[4],
                "spec": r[5], "ver": r[6], "date": r[7], "sec": r[8], "ci": r[9],
                "ct": r[10], "emb": r[11], "path": r[12], "hash": r[13]} for r in rows])


def _existing_hash(db, doc_slug):
    return db.execute(
        text("SELECT content_hash FROM guideline_chunks WHERE doc_slug = :s LIMIT 1"),
        {"s": doc_slug}).scalar()


def _check_db(db) -> bool:
    row = db.execute(text("SELECT current_user, current_database()")).fetchone()
    log.info("Connected as %s @ %s", row[0], row[1])
    tbl = db.execute(text("""
        SELECT 1 FROM pg_tables WHERE schemaname='public' AND tablename='guideline_chunks'
    """)).fetchone()
    if not tbl:
        log.error("guideline_chunks table MISSING — apply db/schema.sql first.")
        return False
    can_ins = db.execute(
        text("SELECT has_table_privilege(current_user,'guideline_chunks','INSERT')")).scalar()
    if not can_ins:
        log.error("User '%s' lacks INSERT on guideline_chunks. Grant it (see schema.sql).", row[0])
        return False
    return True


# ── directory discovery ─────────────────────────────────────────────────────

def _dirs_for_source(settings, source: str):
    """Resolve (dir, source_org) pairs from settings.guideline_dirs, filtered by
    --source. source_org is inferred from the directory name."""
    out = []
    for raw in (settings.guideline_dirs or "").split(","):
        raw = raw.strip()
        if not raw:
            continue
        d = Path(raw)
        name = d.name.lower()
        org = "CAP" if "cap" in name else "ICCR" if "iccr" in name else "CAP"
        if source != "all" and source.upper() != org:
            continue
        if not d.is_dir():
            log.warning("Guideline dir not found: %s", d)
            continue
        out.append((d, org))
    return out


def main():
    ap = argparse.ArgumentParser(description="Embed CAP/ICCR guidelines into guideline_chunks.")
    ap.add_argument("--source", choices=["cap", "iccr", "all"], default="all")
    ap.add_argument("--limit", type=int, default=None, help="max docs to process this run.")
    ap.add_argument("--embed-batch", type=int, default=512,
                    help="chunks buffered/embedded per GPU flush (default 512).")
    ap.add_argument("--reembed", action="store_true",
                    help="re-embed even if content_hash is unchanged.")
    args = ap.parse_args()

    settings = get_settings()

    # Warm up + dim check BEFORE opening the DB (model load can take minutes).
    log.info("Pre-loading embedding model…")
    t0 = time.time()
    dim = len(embed_texts(["guideline embedding warmup"])[0])
    log.info("Embedding model ready: dim=%d (%.1fs)", dim, time.time() - t0)
    if dim != settings.embedding_dim:
        log.error("Model dim (%d) != settings.embedding_dim (%d). Fix config + column together.",
                  dim, settings.embedding_dim)
        sys.exit(1)

    pairs = _dirs_for_source(settings, args.source)
    if not pairs:
        log.error("No guideline directories to process (source=%s, dirs=%r).",
                  args.source, settings.guideline_dirs)
        sys.exit(1)

    db = SessionLocal()
    n_docs = n_skip = n_replace = total_chunks = 0
    try:
        if not _check_db(db):
            sys.exit(1)

        files = []
        for d, org in pairs:
            for p in sorted(d.glob("*.docx")):
                if p.name.startswith(("~$", ".")):
                    continue   # skip Word lock/temp files and hidden files
                files.append((p, org))
        if args.limit:
            files = files[:args.limit]
        log.info("Found %d guideline .docx across %d dir(s).", len(files), len(pairs))

        buf_meta, buf_text = [], []   # meta: (org,slug,title,organ,spec,ver,date,sec,ci,path,hash)

        def flush():
            nonlocal total_chunks
            if not buf_text:
                return
            vectors = embed_texts(buf_text)
            rows = [(m[0], m[1], m[2], m[3], m[4], m[5], m[6], m[7], m[8], m[9], t,
                     _vector_literal(v), m[10], m[11])
                    for m, t, v in zip(buf_meta, buf_text, vectors)]
            _bulk_insert(db, rows)
            db.commit()
            total_chunks += len(rows)
            buf_meta.clear()
            buf_text.clear()

        for path, org in files:
            meta = parse_guideline_filename(path.name, org)
            ex = _extract(path)
            if ex is None:
                continue
            blocks, full_text, title, doc_date = ex
            # Organ from the authoritative inline title (cryptic filenames like
            # 'CXC' don't map to organs); fall back to the filename parse.
            organ = organ_from_title(title, org) or meta["organ"]
            content_hash = hashlib.sha256(full_text.encode("utf-8", "ignore")).hexdigest()

            prior = _existing_hash(db, meta["doc_slug"])
            if prior == content_hash and not args.reembed:
                n_skip += 1
                continue
            if prior is not None:
                # Version/content change → replace this document's rows.
                db.execute(text("DELETE FROM guideline_chunks WHERE doc_slug = :s"),
                           {"s": meta["doc_slug"]})
                db.commit()
                n_replace += 1

            for section, ci, chunk, kind in _doc_chunks(blocks, settings.rag_max_chunk_chars,
                                                        settings.rag_chunk_overlap_chars):
                # Prepend a compact context header so organ/section-specific queries
                # match even when the chunk body doesn't repeat the organ name.
                header = f"[{org} {organ} — {section}]".strip(" —")
                buf_meta.append((org, meta["doc_slug"], kind, title, organ,
                                 meta["specimen_type"], meta["version"], doc_date,
                                 section[:250], ci, str(path), content_hash))
                buf_text.append(f"{header}\n{chunk}")
            n_docs += 1

            if len(buf_text) >= args.embed_batch:
                flush()
                log.info("… %d docs embedded (%d skipped, %d replaced) | %d chunks",
                         n_docs, n_skip, n_replace, total_chunks)

        flush()
        log.info("Done. %d doc(s) embedded into %d chunk(s); %d skipped (unchanged), "
                 "%d replaced (version change).", n_docs, total_chunks, n_skip, n_replace)
    except Exception as e:
        db.rollback()
        log.error("Ingestion failed: %s", e, exc_info=True)
        raise
    finally:
        db.close()


if __name__ == "__main__":
    main()
